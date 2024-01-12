def write_srt(result, srt_path):
    from whisper.utils import get_writer
    writer = get_writer("srt","")
    options = dict()
    options["max_line_width"] = None
    options["max_line_count"] = None
    options["highlight_words"] = False
    try:
        writer(result, srt_path, options)
    except:
        print("Attempting to use an older Version of Whisper")
        writer(result, srt_path)
    print("DONE writing SRT: " + str(srt_path))

def write_txt(text, txt_path):
    try:
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print("DONE writing TXT: " + str(txt_path))
    except Exception as error:
        print("Writing TXT failed:" + str(error))

def add_hyperlink(paragraph, url, text):
    import docx
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def write_docx(speaker_segments:list,translated_segments:list,scriptFilename:str, sourcefile="", translated=False):
    from docx import Document
    from docx.shared import Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from whisper.utils import format_timestamp
    from pathlib import Path
    from docx.shared import RGBColor

    # DIN A4 page setup
    document = Document('utils/TonTexterBase.docx')
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)

    # Setup Styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'

    title = document.styles['Title']
    font = title.font
    font.name = 'Calibri'
    font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    title2 = document.styles['Heading 2']
    font = title2.font
    font.name = 'Calibri'
    font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    # Document header
    heading = document.add_heading('Transkription', 0)
    # File Name
    p = document.add_paragraph('Datei:  ' + Path(sourcefile).name)
    p.style = document.styles['Normal']

    document.add_heading('Revolutioniere deinen Workflow!', 2)
    p = document.add_paragraph('Dieses Transkript wurde mit Ton-Texter generiert. Revolutioniere deinen Video-Editing Workflow mit Text-Based Editing powered by: ')
    add_hyperlink(p,'https://ton-texter.de','Ton-Texter')

    latest_index = 0

    # initialize table
    table = document.add_table(rows=1, cols=4)
    table.style = document.styles['Table Grid']
    hdr_cells = table.rows[0].cells

    paragraph = hdr_cells[0].paragraphs[0]
    run = paragraph.add_run('ID')
    run.bold = True
    hdr_cells[0].width = Mm(9.4)

    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph = hdr_cells[1].paragraphs[0]
    run = paragraph.add_run('Timecode')
    run.bold = True
    hdr_cells[1].width = Mm(30)

    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph = hdr_cells[2].paragraphs[0]
    run = paragraph.add_run('Text')
    hdr_cells[2].width = Mm(50)
    run.bold = True

    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph = hdr_cells[3].paragraphs[0]
    run = paragraph.add_run('Translation')
    hdr_cells[3].width = Mm(50)
    run.bold = True
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER

    # add the speaker segments to the table
    for index_speaker,speaker in enumerate(speaker_segments):

        if len(speaker.segments) != 0:

            speaker_name = table.add_row()
            run = speaker_name.cells[0].paragraphs[0].add_run(speaker.speaker)
            speaker_name.cells[0].merge(speaker_name.cells[3])
            speaker_name.cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True


            trans_text = ""
            translation = translated_segments[index_speaker]
            start = 0
            end = 0

            for index, s in enumerate(speaker.segments):

                start = format_timestamp(s['start'],True,':')

                end = format_timestamp(s['end'],True,':')

                row_cells = table.add_row().cells

                row_cells[0].text = str(latest_index + index + 1)
                row_cells[0].width = Mm(9.4)
                row_cells[1].text = start + " - " + end
                row_cells[2].text = s['text']

                # When a translation is available add it into the table
                if translated:
                    trans_text = translation.segments[index]['text']

                row_cells[3].text = trans_text

                trans_text = ""

            latest_index += len(speaker.segments)

        

    paragraph = document.add_paragraph()
    run = paragraph.add_run('Dieses Transkript wurde automatisch durch Ton-Texter generiert und benötigt dementsprechende Überarbeitung. Stellen Sie sicher, dass die Inhalte mit der originalen Audiospur abgeglichen werden.')
    run.italic = True   

    document.save(scriptFilename)
    print("DONE writing DOCX: " + str(scriptFilename))